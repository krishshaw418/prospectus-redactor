from docx.document import Document
from docx.table import Table, _Cell
from model import ParagraphRef, Location

def iter_paragraphs(parent):
    """
        Recursively yield paragraphs from every part of the document:\n
    """

    if isinstance(parent, Document):

        for paragraph in parent.paragraphs:
            yield ParagraphRef(paragraph, Location.BODY)

        for table in parent.tables:
            yield from iter_paragraphs(table)

    elif isinstance(parent, Table):

        for row in parent.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)

    elif isinstance(parent, _Cell):

        for paragraph in parent.paragraphs:
            yield ParagraphRef(paragraph, Location.TABLE)

        for table in parent.tables:
            yield from iter_paragraphs(table)